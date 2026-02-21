"""Read sample_data.xlsx and Template for Lable.docx and print their contents."""
import openpyxl
import sys

# ── Excel ──────────────────────────────────────────────────────────────────────
print("=" * 60)
print("SAMPLE DATA EXCEL")
print("=" * 60)
wb = openpyxl.load_workbook(r'Samples\sample_data.xlsx')
ws = wb.active
rows = list(ws.iter_rows(values_only=True))
if rows:
    print("Columns:")
    for i, col in enumerate(rows[0]):
        print(f"  [{i}] {col}")
    print(f"\nTotal data rows: {len(rows)-1}")
    print("\nFirst 3 rows:")
    for row in rows[1:4]:
        print(" ", row)

# ── DOCX ───────────────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("TEMPLATE DOCX CONTENT")
print("=" * 60)
try:
    from docx import Document
    doc = Document(r'Samples\Template for Lable.docx')
    for para in doc.paragraphs:
        if para.text.strip():
            print(repr(para.text))
    for table in doc.tables:
        print("\n[TABLE]")
        for row in table.rows:
            print([cell.text.strip() for cell in row.cells])
except ImportError:
    print("python-docx not installed, trying raw XML...")
    import zipfile, re
    with zipfile.ZipFile(r'Samples\Template for Lable.docx') as z:
        with z.open('word/document.xml') as f:
            content = f.read().decode('utf-8')
        text = re.sub(r'<[^>]+>', ' ', content)
        text = re.sub(r'\s+', ' ', text).strip()
        print(text[:3000])
